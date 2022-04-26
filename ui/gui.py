import sys
import tkinter as tk
import subprocess
import os
import win32print
import re

from PyQt5.QtCore import Qt
from PyQt5.QtGui import QStandardItemModel, QStandardItem, QIcon
from PyQt5.QtWidgets import QWidget, QDesktopWidget, QLabel, QGridLayout, QLineEdit, QTreeView, \
    QPushButton, QHBoxLayout, QCheckBox, \
    QMessageBox
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.shared import Inches, Cm, Pt
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Side, Border
from openpyxl.worksheet.page import PageMargins

from common import enum
from do.obj_info import ObjInfo
from ui.table_r import TableR
from utils import dir_util
from tkinter import filedialog
from docx2pdf import convert


class MainWin(QWidget):
    WINDOW_TITLE = '标准目录生成工具'
    WINDOW_WIDTH = 1000
    WINDOW_HEIGHT = 600

    LEFT_HEADER = ['目录\\文件']
    RIGHT_HEADER = ['',
                    '题名',
                    '页次',
                    '',
                    # '页码(test)'
                    ]

    ROOT_LINE = 0

    DOCX_TOP_TITLE = ["序号", "题名", "页次"]
    DOCX_TOP_WIDTH = [1, 8, 1]

    def __init__(self, app):
        super().__init__()
        self.app = app
        self.in_lab = QLabel('输入目录：')
        self.in_edit = QLineEdit()
        self.in_edit.setReadOnly(True)
        self.in_btn = QPushButton('选择文件')
        self.icon_map = {}

        self.left_lab = QLabel('目录列表')
        self.left_tree = QTreeView()
        self.left_model = QStandardItemModel()

        self.right_lab = QLabel('题名信息')
        self.right_tab = TableR()

        self.new_btn = QPushButton('新增')
        self.top_btn = QPushButton('上移')
        self.down_btn = QPushButton('下移')
        self.del_btn = QPushButton('删除')

        self.selection_lab = QLabel('输出文件：')
        self.word_check = QCheckBox('Word')
        self.excel_check = QCheckBox('Excel')

        self.out_lab = QLabel('输出目录：')
        self.out_edit = QLineEdit()
        self.out_btn = QPushButton('选择目录')

        self.creat_btn = QPushButton('生成')
        self.print_btn = QPushButton('一键打印')
        self.exit_btn = QPushButton('结束')
        self.init_ui()

    def init_ui(self):
        self.resize(self.WINDOW_WIDTH, self.WINDOW_HEIGHT)
        # 固定大小
        # self.setFixedSize(self.WINDOW_WIDTH, self.WINDOW_HEIGHT)
        # 禁用最大最小
        # self.setWindowFlags(Qt.WindowType.WindowCloseButtonHint)
        self.setWindowFlags(Qt.WindowType.Window)
        # 居中
        self.center()
        self.setWindowTitle(self.WINDOW_TITLE)
        win_icon = QIcon()
        win_icon.addFile(r"static\icon\win.png")
        self.setWindowIcon(win_icon)
        # 内容
        grid = QGridLayout()
        grid.setSpacing(20)
        grid.addLayout(self.init_in_model(), 0, 0)
        grid.addLayout(self.init_dir_model(), 1, 0)
        grid.addLayout(self.init_selection_model(), 2, 0)
        grid.addLayout(self.init_out_model(), 3, 0)
        grid.addLayout(self.init_final_model(), 4, 0)

        self.setLayout(grid)
        self.show()

    def init_in_model(self):
        grid = QGridLayout()
        grid.setSpacing(10)
        grid.addWidget(self.in_lab, 0, 0)
        grid.addWidget(self.in_edit, 0, 1)
        grid.addWidget(self.in_btn, 0, 2)

        self.in_btn.clicked.connect(self.select_source)
        return grid

    def init_dir_model(self):
        grid = QGridLayout()
        grid.setSpacing(10)
        # 将左边添加到容器中
        grid_l = QGridLayout()
        grid_l.addWidget(self.left_lab, 0, 0)
        grid_l.addWidget(self.left_tree, 1, 0)
        grid.addLayout(grid_l, 0, 0)
        self.left_tree.setModel(self.left_model)
        # 添加右边
        grid_r_btn = QGridLayout()
        grid_r_btn.addWidget(self.new_btn, 0, 0)
        self.new_btn.clicked.connect(self.right_tab.new)
        grid_r_btn.addWidget(self.top_btn, 0, 1)
        self.top_btn.clicked.connect(self.right_tab.top)
        grid_r_btn.addWidget(self.down_btn, 0, 2)
        self.down_btn.clicked.connect(self.right_tab.down)
        grid_r_btn.addWidget(self.del_btn, 0, 3)
        self.del_btn.clicked.connect(self.right_tab.remove)
        grid_r = QGridLayout()
        grid_r.addWidget(self.right_lab, 0, 0)
        grid_r.addLayout(grid_r_btn, 1, 0)
        grid_r.addWidget(self.right_tab, 2, 0)
        grid.addLayout(grid_r, 0, 1)
        self.init_l_r()
        return grid

    def init_l_r(self):
        # 左
        self.left_model.clear()
        self.left_model.setHorizontalHeaderLabels(self.LEFT_HEADER)
        # 右
        self.right_tab.clearContents()
        self.right_tab.setRowCount(0)
        self.right_tab.setHorizontalHeaderLabels(self.RIGHT_HEADER)
        self.right_tab.init_ui()
        self.ROOT_LINE = 0

    def update_dir_model(self, root_path):
        self.init_l_r()
        # 用来存放所有的目录路径
        file_list = []
        root_info = dir_util.foreach_dir(root_path, None, file_list)
        self.foreach_dir(root_info.dir_list)
        self.foreach_file(root_info.file_list)

    def foreach_dir(self, obj_list: [ObjInfo], parent_item=None):
        for line, obj_info in enumerate(obj_list):
            now_item_l = self.add_item_l(obj_info.obj_name, obj_info, enum.DIR_TREE_ROLE, parent_item)
            if parent_item is None:
                self.right_tab.append_row(self.ROOT_LINE, obj_info, enum.DIR_TREE_ROLE)
                self.ROOT_LINE += 1
            if obj_info.dir_list is not None:
                self.foreach_dir(obj_info.dir_list, now_item_l)
            if obj_info.file_list is not None:
                for child_file in obj_info.file_list:
                    if child_file is not None:
                        self.add_item_l(child_file.obj_name, child_file, enum.FILE_TREE_ROLE, now_item_l)

    def foreach_file(self, obj_list: [ObjInfo]):
        for line, obj_info in enumerate(obj_list):
            self.add_item_l(obj_info.obj_name, obj_info, enum.FILE_TREE_ROLE)
            self.right_tab.append_row(self.ROOT_LINE, obj_info, enum.FILE_TREE_ROLE)
            self.ROOT_LINE += 1

    def add_item_l(self, obj_name, obj_data: ObjInfo = None, role=None, parent_item=None):
        item_l = QStandardItem()
        item_l.setText(obj_name)
        item_l.setData(obj_data, role)
        item_l.setEditable(False)
        if parent_item is None:
            self.left_model.appendRow(item_l)
        else:
            parent_item.appendRow([item_l])
        # 图标
        dir_icon = QIcon()
        if role == enum.DIR_TREE_ROLE:
            dir_icon.addFile(r"static\icon\folder.ico")
        else:
            file_type = obj_data.obj_type
            if file_type is enum.IMG_TYPE:
                dir_icon.addFile(r"static\icon\jpg.png")
            elif file_type is enum.DOC_TYPE:
                dir_icon.addFile(r"static\icon\doc.png")
            elif file_type is enum.PDF_TYPE:
                dir_icon.addFile(r"static\icon\pdf.png")
            elif file_type is enum.XLS_TYPE:
                dir_icon.addFile(r"static\icon\xls.png")
            elif file_type is enum.OFD_TYPE:
                dir_icon.addFile(r"static\icon\ofd.png")
            elif file_type is enum.PPT_TYPE:
                dir_icon.addFile(r"static\icon\ppt.png")
            else:
                dir_icon.addFile(r"static\icon\pictures.ico")
        item_l.setIcon(dir_icon)
        return item_l

    def init_selection_model(self):
        grid = QHBoxLayout()
        grid.setSpacing(10)
        grid.addWidget(self.selection_lab, 0, Qt.AlignmentFlag.AlignLeft)
        self.selection_lab.setFixedWidth(60)
        grid.addWidget(self.word_check, 0, Qt.AlignmentFlag.AlignLeft)
        self.word_check.setFixedWidth(60)
        self.word_check.setChecked(True)
        grid.addWidget(self.excel_check, 15, Qt.AlignmentFlag.AlignLeft)
        self.excel_check.setFixedWidth(60)
        return grid

    def init_out_model(self):
        grid = QGridLayout()
        grid.setSpacing(10)
        grid.addWidget(self.out_lab, 0, 0)
        grid.addWidget(self.out_edit, 0, 1)
        grid.addWidget(self.out_btn, 0, 3)
        self.out_edit.setReadOnly(True)
        self.out_btn.clicked.connect(self.select_target)
        return grid

    def init_final_model(self):
        grid = QHBoxLayout()
        grid.setSpacing(10)
        grid.addWidget(self.creat_btn, 10, Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.print_btn, 0, Qt.AlignmentFlag.AlignRight)
        grid.addWidget(self.exit_btn, 0, Qt.AlignmentFlag.AlignRight)
        self.creat_btn.clicked.connect(self.create_file)
        self.print_btn.clicked.connect(self.print_file)
        self.exit_btn.clicked.connect(self.exit_window)
        return grid

    # 控制窗口显示在屏幕中心的方法
    def center(self):
        # 获得窗口
        qr = self.frameGeometry()
        # 获得屏幕中心点
        cp = QDesktopWidget().availableGeometry().center()
        # 显示到屏幕中心
        qr.moveCenter(cp)
        self.move(qr.topLeft())

    def select_source(self):
        root = tk.Tk()
        root.withdraw()
        root_path = filedialog.askdirectory()
        root_path = root_path.replace('/', '\\')
        self.in_edit.setText(root_path)
        if root_path is not None and root_path != '':
            self.update_dir_model(root_path)

    def select_target(self):
        root = tk.Tk()
        root.withdraw()
        root_path = filedialog.askdirectory()
        self.out_edit.setText(root_path)

    def create_file(self):
        try:
            if self.right_tab.rowCount() <= 0:
                QMessageBox.critical(self, "系统提示", "目录为空！")
                return
            if self.word_check.isChecked() is False and self.excel_check.isChecked() is False:
                QMessageBox.critical(self, "系统提示", "请选择输出文件类型！")
                return
            # 存储位置
            save_path = self.out_edit.text()
            if save_path is None or save_path == "":
                save_path = self.in_edit.text()
            file_name_array = re.split(r'[\\/]', save_path)
            file_name = file_name_array[file_name_array.__len__() - 1]
            file_path = save_path + "\\" + file_name + "目录"
            if self.word_check.isChecked():
                docx_file_path = file_path + ".docx"
                self.create_doc(docx_file_path)
            if self.excel_check.isChecked():
                xlsx_file_path = file_path + ".xlsx"
                self.create_xls(xlsx_file_path)
            QMessageBox.information(self, "系统提示", "目录文件生成成功！", QMessageBox.Yes)
        except Exception as e:
            print("引发异常：", repr(e))
            QMessageBox.critical(self, "系统异常", repr(e))

    def print_file(self):
        try:
            if self.right_tab.rowCount() <= 0:
                QMessageBox.critical(self, "系统提示", "目录为空！")
                return
            if sys.platform == 'win32':
                # 存储位置
                save_path = self.out_edit.text()
                if save_path is None or save_path == "":
                    save_path = self.in_edit.text()
                file_name = "~" + re.split(r'[\\/]', save_path)[1]
                docx_file_path = save_path + "\\" + file_name + "temp.docx"
                pdf_file_path = save_path + "\\" + file_name + "temp.pdf"
                self.create_doc(docx_file_path)
                # docx_file_path = docx_file_path.replace("\\", "/")
                pdf_file = open(pdf_file_path, 'w')
                pdf_file.close()
                convert(docx_file_path, pdf_file_path)
                printer = win32print.GetDefaultPrinter()
                args = [r"static\PDFtoPrinter.exe",
                        f"{pdf_file_path}",
                        f"{printer}",
                        ]
                subprocess.run(args, encoding="utf-8", shell=True)
                # 删除临时文件
                os.remove(docx_file_path)
                os.remove(pdf_file_path)
                QMessageBox.information(self, "系统提示", f"已发送至打印机:{printer}！", QMessageBox.Yes)
        except Exception as e:
            print("引发异常：", repr(e))
            QMessageBox.critical(self, "系统异常", repr(e))

    def exit_window(self):
        sys.exit(self.app.exec_())

    def create_doc(self, docx_file_path):
        doc = Document("./static/template/template-1.docx")
        # 删除第一行
        doc._body.clear_content()
        # 定义表格
        rows = 2
        cols = 3
        style = 'Table Grid'
        table = doc.add_table(rows, cols, style)
        table.autofit = True
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # 标题
        table.cell(0, 0).merge(table.cell(0, 2))
        table.rows[0].height = Cm(2)
        title_cell = table.cell(0, 0)
        # 垂直居中
        title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        title_cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        title_run = title_cell.paragraphs[0].add_run("卷宗目录")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.name = '宋体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 表头
        colums = table.column_cells(0)
        top_cells = table.row_cells(1)
        table.rows[1].height = Cm(1.2)
        for i in range(top_cells.__len__()):
            top_cell = top_cells[i]
            top_cell.width = Inches(self.DOCX_TOP_WIDTH[i])
            top_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            top_cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            top_cell_run = top_cell.paragraphs[0].add_run(self.DOCX_TOP_TITLE[i])
            top_cell_run.font.size = Pt(11)
            top_cell_run.font.bold = True
            top_cell_run.font.name = '宋体'
            top_cell_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        ignore_count = 0
        # 内容
        for i in range(self.right_tab.rowCount()):
            data_item = self.right_tab.item(i, 1)
            page_item = self.right_tab.item(i, 2)
            if data_item.data(enum.IGNORE_ROLE) is True:
                ignore_count += 1
                continue
            # 新增一行
            table.add_row()
            i -= ignore_count
            body_rows_cells = table.row_cells(i + 2)
            table.rows[i + 2].height = Cm(1.2)
            # 序号 题名 页次
            index = i + 1
            name = data_item.text()
            page = page_item.text()
            # 序号
            body_rows_cell_0 = body_rows_cells[0]
            body_rows_cell_0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            body_rows_cell_0.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            body_rows_cell_0.width = Inches(self.DOCX_TOP_WIDTH[0])
            body_rows_cell_run_0 = body_rows_cell_0.paragraphs[0].add_run(str(index))
            body_rows_cell_run_0.font.size = Pt(11)
            body_rows_cell_run_0.font.name = '宋体'
            body_rows_cell_run_0._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 题名
            body_rows_cell_1 = body_rows_cells[1]
            body_rows_cell_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            body_rows_cell_1.width = Inches(self.DOCX_TOP_WIDTH[1])
            body_rows_cell_run_1 = body_rows_cell_1.paragraphs[0].add_run(name)
            body_rows_cell_run_1.font.size = Pt(11)
            body_rows_cell_run_1.font.name = '宋体'
            body_rows_cell_run_1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            # 页次
            body_rows_cell_2 = body_rows_cells[2]
            body_rows_cell_2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            body_rows_cell_2.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            body_rows_cell_2.width = Inches(self.DOCX_TOP_WIDTH[2])
            body_rows_cells_run_2 = body_rows_cell_2.paragraphs[0].add_run(page)
            body_rows_cells_run_2.font.size = Pt(11)
            body_rows_cells_run_2.font.name = '宋体'
            body_rows_cells_run_2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.save(docx_file_path)

    def create_xls(self, xlsx_file_path):
        wb = Workbook()
        # 边框
        # 边框样式，可选dashDot、dashDotDot、dashed、dotted、double、hair、medium、mediumDashDot、mediumDashDotDot、mediumDashed、slantDashDot、thick、thin
        side = Side(style="thin")
        border = Border(
            top=side,  # 上
            bottom=side,  # 下
            left=side,  # 左
            right=side,  # 右
        )
        # 列宽
        work_sheet = wb.active
        work_sheet.page_margins = PageMargins(left=Cm(2.8).inches, right=Cm(2.6).inches, top=Cm(3.7).inches, bottom=Cm(3.5).inches, header=0, footer=0)
        work_sheet.column_dimensions["A"].width = 8
        work_sheet.column_dimensions["B"].width = 60
        work_sheet.column_dimensions["C"].width = 8
        # 标题
        title_font = Font(
            name="宋体",  # 字体
            size=16,  # 字体大小
            bold=True,  # 是否加粗，True/False
        )
        title_align = Alignment(horizontal='center', vertical='center')
        work_sheet.append(["卷宗目录"])
        work_sheet.row_dimensions[1].height = 57
        work_sheet.merge_cells(range_string='A1:C1')
        work_sheet["A1"].font = title_font
        work_sheet["A1"].alignment = title_align
        work_sheet["A1"].border = border
        work_sheet["B1"].border = border
        work_sheet["C1"].border = border
        # 表头
        top_font = Font(
            name="宋体",  # 字体
            size=11,  # 字体大小
            bold=True  # 是否加粗，True/False
        )
        top_align = Alignment(horizontal='center', vertical='center')
        work_sheet.append(["序号", "题名", "目录"])
        work_sheet.row_dimensions[2].height = 34.5
        work_sheet["A2"].font = top_font
        work_sheet["B2"].font = top_font
        work_sheet["C2"].font = top_font
        work_sheet["A2"].alignment = top_align
        work_sheet["B2"].alignment = top_align
        work_sheet["C2"].alignment = top_align
        work_sheet["A2"].border = border
        work_sheet["B2"].border = border
        work_sheet["C2"].border = border
        # 内容
        body_font = Font(
            name="宋体",  # 字体
            size=11  # 字体大小
        )
        body_align_1 = Alignment(horizontal='center', vertical='center')
        body_align_2 = Alignment(vertical='center', wrapText=True)
        ignore_count = 0
        for i in range(self.right_tab.rowCount()):
            data_item = self.right_tab.item(i, 1)
            page_item = self.right_tab.item(i, 2)
            if data_item.data(enum.IGNORE_ROLE) is True:
                ignore_count += 1
                continue
            i -= ignore_count
            index = i + 1
            name = data_item.text()
            page = page_item.text()
            work_sheet.append([str(index), name, page])
            index = index + 2
            work_sheet.row_dimensions[index].height = 34.5
            work_sheet["A" + str(index)].font = body_font
            work_sheet["B" + str(index)].font = body_font
            work_sheet["C" + str(index)].font = body_font
            work_sheet["A" + str(index)].alignment = body_align_1
            work_sheet["B" + str(index)].alignment = body_align_2
            work_sheet["C" + str(index)].alignment = body_align_1
            work_sheet["A" + str(index)].border = border
            work_sheet["B" + str(index)].border = border
            work_sheet["C" + str(index)].border = border
        # 存储
        wb.save(xlsx_file_path)























