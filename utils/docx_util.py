from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import uuid


def creat_docx(file_path):
    doc = Document()
    write_title(doc, "目   录")
    write_dir_page(doc, "1.这是一个段落这是一个段落asdasdasd这是一个段落这是一个段落这是一个段落这是一个段落一个段落是一个段落这是一个段落这是一个段落这是一个段落这是一个段落这是一个段落这是一个段落这是一个段落这是一个段落", "112")
    doc.save(file_path)


def write_title(doc: Document, title):
    paragraph = doc.add_paragraph()
    paragraph.add_run(title)
    for run in paragraph.runs:
        # run.font.bold = True
        run.font.size = Pt(18)
        run.font.name = "方正小标宋简体"
        r = run.element.rPr.rFonts
        r.set(qn("w:eastAsia"), "方正小标宋简体")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def write_paragraph(doc: Document, context):
    paragraph = doc.add_paragraph()
    paragraph.add_run(context)
    for run in paragraph.runs:
        # run.font.bold = True
        run.font.size = Pt(14)
        run.font.name = "宋体"
        r = run.element.rPr.rFonts
        r.set(qn("w:eastAsia"), "宋体")


def write_dir_page(doc: Document, dir_name, page):
    temp = 30-len(dir_name + page) % 30
    context = dir_name
    for i in range(temp):
        context += '..'
    context += page
    write_paragraph(doc, context)


if __name__ == '__main__':
    file = "../" + str(uuid.uuid4()) + ".docx"
    creat_docx(file)
